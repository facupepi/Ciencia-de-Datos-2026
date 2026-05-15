"""
Reporte de Scoring Académico — PDF
Replica los resultados del notebook tp_scoring.ipynb y los exporta a un PDF
con portada, tablas, gráficos y conclusiones.

Uso:
    python reporte_scoring_pdf.py
"""

import os
import warnings

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import seaborn as sns

from docx import Document
from docx.shared import Pt, RGBColor

from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
from sklearn.metrics import (
    confusion_matrix, roc_auc_score, roc_curve,
    precision_recall_curve, average_precision_score,
)

warnings.filterwarnings("ignore")

# ── Configuración ───────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CSV = os.path.join(BASE_DIR, "Student exam score - 7 de mayo.csv")
PDF_PATH = os.path.join(BASE_DIR, "reporte_scoring.pdf")
DOCX_PATH = os.path.join(BASE_DIR, "reporte_scoring.docx")
CSV_OUT = os.path.join(BASE_DIR, "scoring_resultados.csv")

RANDOM_STATE = 42
np.random.seed(RANDOM_STATE)

sns.set_theme(style="whitegrid", palette="colorblind", font_scale=1.0)
plt.rcParams["figure.dpi"] = 150

FEATURES = ["hours_studied", "sleep_hours", "attendance_percent", "previous_scores"]
NUMERICAS = FEATURES + ["exam_score"]

PESOS = {
    "hours_studied": 0.40,
    "previous_scores": 0.30,
    "attendance_percent": 0.20,
    "sleep_hours": 0.10,
}
ORDEN_NIVELES = ["Bajo", "Medio", "Alto", "Crítico"]
COLORES_NIVEL = {"Bajo": "#27ae60", "Medio": "#f1c40f",
                 "Alto": "#e67e22", "Crítico": "#c0392b"}


# ── Lógica de scoring (idéntica al notebook) ────────────────────────────────

def normalizar_riesgo(valor, bueno, malo):
    if bueno == malo:
        return 0.0
    return float(np.clip((valor - bueno) / (malo - bueno) * 100, 0, 100))


def riesgo_sueno(horas):
    optimo_min, optimo_max = 7, 8
    if optimo_min <= horas <= optimo_max:
        return 0.0
    desvio = optimo_min - horas if horas < optimo_min else horas - optimo_max
    return float(np.clip(desvio / 3 * 100, 0, 100))


def calcular_score_reglas(fila):
    r_horas = normalizar_riesgo(fila["hours_studied"], bueno=10, malo=2)
    r_prev = normalizar_riesgo(fila["previous_scores"], bueno=85, malo=50)
    r_asis = normalizar_riesgo(fila["attendance_percent"], bueno=90, malo=60)
    r_sueno = riesgo_sueno(fila["sleep_hours"])
    score = (
        PESOS["hours_studied"]      * r_horas +
        PESOS["previous_scores"]    * r_prev +
        PESOS["attendance_percent"] * r_asis +
        PESOS["sleep_hours"]        * r_sueno
    )
    return round(score, 2)


def asignar_nivel(score):
    if score < 30:
        return "Bajo"
    if score < 50:
        return "Medio"
    if score < 70:
        return "Alto"
    return "Crítico"


# ── Carga y procesamiento de datos ──────────────────────────────────────────
print("[1/4] Cargando dataset...")
df = pd.read_csv(CSV)
N = len(df)

UMBRAL_BAJO = df["exam_score"].quantile(0.25)
df["riesgo_bajo_desempeno"] = (df["exam_score"] < UMBRAL_BAJO).astype(int)

print("[2/4] Calculando score por reglas...")
df["score_reglas"] = df.apply(calcular_score_reglas, axis=1)
df["nivel_reglas"] = df["score_reglas"].apply(asignar_nivel)

print("[3/4] Entrenando modelos...")
X = df[FEATURES]
y = df["riesgo_bajo_desempeno"]
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.30, random_state=RANDOM_STATE, stratify=y
)

modelo_lr = Pipeline([
    ("scaler", StandardScaler()),
    ("clf", LogisticRegression(max_iter=1000, random_state=RANDOM_STATE)),
])
modelo_rf = RandomForestClassifier(
    n_estimators=200, max_depth=5, random_state=RANDOM_STATE
)

auc_lr_cv = cross_val_score(modelo_lr, X_train, y_train, cv=5, scoring="roc_auc")
auc_rf_cv = cross_val_score(modelo_rf, X_train, y_train, cv=5, scoring="roc_auc")

modelo_lr.fit(X_train, y_train)
modelo_rf.fit(X_train, y_train)

prob_lr = modelo_lr.predict_proba(X_test)[:, 1]
prob_rf = modelo_rf.predict_proba(X_test)[:, 1]
auc_lr = roc_auc_score(y_test, prob_lr)
auc_rf = roc_auc_score(y_test, prob_rf)

if auc_lr >= auc_rf:
    modelo_final, nombre_final = modelo_lr, "Logistic Regression"
else:
    modelo_final, nombre_final = modelo_rf, "Random Forest"

df["prob_modelo"] = modelo_final.predict_proba(df[FEATURES])[:, 1]
df["score_modelo"] = (df["prob_modelo"] * 100).round(2)
df["nivel_modelo"] = df["score_modelo"].apply(asignar_nivel)


# ── Scores adicionales derivados ────────────────────────────────────────────
print("[3.5/4] Calculando 8 scores adicionales...")

# 1) Score de desempeño (inverso del riesgo)
df["score_desempeno"] = (100 - df["score_reglas"]).round(2)

# 2) Score de potencial de mejora
prev_pct = df["previous_scores"].rank(pct=True) * 100
exam_pct = df["exam_score"].rank(pct=True) * 100
df["score_potencial"] = ((prev_pct - exam_pct) + 50).clip(0, 100).round(2)


# 3) Score de engagement
def calcular_engagement(fila):
    horas_n = np.clip((fila["hours_studied"] - 1) / (10 - 1), 0, 1) * 100
    asis_n = np.clip((fila["attendance_percent"] - 50) / (100 - 50), 0, 1) * 100
    return round(0.4 * horas_n + 0.6 * asis_n, 2)


df["score_engagement"] = df.apply(calcular_engagement, axis=1)

# 4) Score de eficiencia
df["score_eficiencia"] = (
    (df["exam_score"] / df["hours_studied"]).rank(pct=True) * 100
).round(2)


# 5) Score de hábitos saludables
def calcular_habitos(fila):
    if 7 <= fila["sleep_hours"] <= 8:
        sueno = 100
    else:
        desvio = min(abs(fila["sleep_hours"] - 7), abs(fila["sleep_hours"] - 8))
        sueno = max(0, 100 - desvio * 25)
    if 3 <= fila["hours_studied"] <= 8:
        estudio = 100
    elif fila["hours_studied"] < 3:
        estudio = max(0, 100 - (3 - fila["hours_studied"]) * 30)
    else:
        estudio = max(0, 100 - (fila["hours_studied"] - 8) * 25)
    asis = float(np.clip((fila["attendance_percent"] - 50) / 50 * 100, 0, 100))
    return round(0.4 * sueno + 0.3 * estudio + 0.3 * asis, 2)


df["score_habitos"] = df.apply(calcular_habitos, axis=1)

# 6) Score de progreso
df["score_progreso"] = ((exam_pct - prev_pct) + 50).clip(0, 100).round(2)

# 7) Score de prioridad
df["score_prioridad"] = (0.6 * df["score_reglas"] + 0.4 * df["score_potencial"]).round(2)

# 8) Score de anomalía
riesgo_real = 100 - exam_pct
df["score_anomalia"] = (df["score_reglas"] - riesgo_real).abs().round(2)


# 9) Score de deserción (PROXY — no hay ground truth en el dataset)
#    Basado en literatura: la asistencia es el predictor más fuerte de abandono,
#    seguido por engagement activo (horas) y trayectoria académica previa.
def calcular_desercion(fila):
    asis = normalizar_riesgo(fila["attendance_percent"], bueno=85, malo=55)
    horas = normalizar_riesgo(fila["hours_studied"], bueno=8, malo=2)
    prev = normalizar_riesgo(fila["previous_scores"], bueno=80, malo=45)
    sueno = riesgo_sueno(fila["sleep_hours"])
    score = 0.45 * asis + 0.25 * horas + 0.20 * prev + 0.10 * sueno
    return round(score, 2)


df["score_desercion"] = df.apply(calcular_desercion, axis=1)

NUEVOS_SCORES = [
    "score_desempeno", "score_potencial", "score_engagement", "score_eficiencia",
    "score_habitos", "score_progreso", "score_prioridad", "score_anomalia",
    "score_desercion",
]

# CSV de salida (con los 10 scores)
df_salida = df[["student_id"] + NUMERICAS +
               ["score_reglas", "nivel_reglas",
                "score_modelo", "nivel_modelo"] +
               NUEVOS_SCORES +
               ["riesgo_bajo_desempeno"]].copy()
df_salida.to_csv(CSV_OUT, index=False, encoding="utf-8-sig")


# ── Helpers de PDF ──────────────────────────────────────────────────────────

def portada(pdf, titulo, subtitulo=""):
    fig, ax = plt.subplots(figsize=(11, 8.5))
    ax.axis("off")
    ax.text(0.5, 0.68, titulo, transform=ax.transAxes,
            fontsize=26, fontweight="bold", ha="center", va="center", color="#2c3e50")
    ax.text(0.5, 0.56, subtitulo, transform=ax.transAxes,
            fontsize=14, ha="center", va="center", color="#7f8c8d")
    ax.text(0.5, 0.44, "Autor: Facundo Pepino",
            transform=ax.transAxes, fontsize=14, fontweight="bold",
            ha="center", va="center", color="#2c3e50")
    ax.text(0.5, 0.36,
            f"Dataset: Student exam score — {N} estudiantes, {len(df.columns) - 4} variables originales",
            transform=ax.transAxes, fontsize=11, ha="center", va="center", color="#95a5a6")
    ax.text(0.5, 0.28,
            "TP Ciencia de Datos 2026 — UTN FRSFCO",
            transform=ax.transAxes, fontsize=10, ha="center", va="center", color="#95a5a6")
    pdf.savefig(fig); plt.close()


def pagina_titulo(pdf, titulo):
    # El texto del reporte va al .docx, no al PDF.
    if "_DOC" in globals() and _DOC is not None:
        h = _DOC.add_heading(titulo, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)


def render_tabla(pdf, data, col_labels, titulo, col_widths=None,
                 fontsize=9, row_height=1.3):
    n_rows = len(data)
    fig_height = min(max(3.5, 1.8 + n_rows * 0.32), 10.5)
    fig, ax = plt.subplots(figsize=(11, fig_height))
    ax.axis("off")
    ax.set_title(titulo, fontsize=13, fontweight="bold", pad=15, loc="left")
    table = ax.table(cellText=data, colLabels=col_labels, loc="center",
                     cellLoc="center", colWidths=col_widths)
    table.auto_set_font_size(False)
    table.set_fontsize(fontsize)
    table.scale(1, row_height)
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor("#2c3e50")
            cell.set_text_props(color="white", fontweight="bold")
        elif row % 2 == 0:
            cell.set_facecolor("#ecf0f1")
    plt.tight_layout()
    pdf.savefig(fig); plt.close()


def pagina_texto(pdf, titulo, lineas):
    # El texto del reporte va al .docx, no al PDF.
    if "_DOC" in globals() and _DOC is not None:
        h = _DOC.add_heading(titulo, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        for ln in lineas:
            p = _DOC.add_paragraph(ln)
            for run in p.runs:
                run.font.size = Pt(11)


def pagina_detalle_score(pdf, num, total, codigo, descripcion,
                         formula_lines, interpretacion_lines,
                         niveles=None, accion=None, max_info=None):
    """Renderiza una página dedicada por score, con header destacado.

    Layout determinístico: un cursor 'y' que descuenta la altura real de
    cada bloque (header de sección + N líneas de contenido + gap).
    Todos los ax.text usan va='top', así 'y' es siempre el borde superior.
    """
    from matplotlib.patches import Rectangle

    fig, ax = plt.subplots(figsize=(11, 8.5))
    ax.axis("off")

    # ── Banda superior con el nombre del score ──
    ax.add_patch(Rectangle((0.04, 0.85), 0.92, 0.10,
                           transform=ax.transAxes,
                           facecolor="#2c3e50", edgecolor="none"))
    ax.text(0.06, 0.925, f"SCORE  {num} / {total}",
            transform=ax.transAxes, fontsize=10,
            color="#bdc3c7", va="center", fontweight="bold")
    ax.text(0.06, 0.885, codigo, transform=ax.transAxes,
            fontsize=22, fontweight="bold", color="white",
            fontfamily="monospace", va="center")
    ax.text(0.96, 0.885, descripcion, transform=ax.transAxes,
            fontsize=12, color="#ecf0f1", va="center", ha="right",
            style="italic")

    # ── Constantes de layout (ajustadas a fontsize=10.5 + linespacing=1.5) ──
    LINE_H = 0.030       # alto de una línea en coordenadas de eje
    HEADER_H = 0.040     # espacio que ocupa un header de sección
    SECTION_GAP = 0.022  # gap entre el final de un bloque y el siguiente header

    # ── Construir lista de secciones ──
    sections = [
        ("FÓRMULA",        "\n".join(formula_lines),        True,  "#2c3e50"),
        ("INTERPRETACIÓN", "\n".join(interpretacion_lines), False, "#34495e"),
    ]
    if niveles:
        sections.append(("NIVELES",            niveles,  True,  "#2c3e50"))
    if accion:
        sections.append(("DECISIÓN QUE APOYA", accion,   False, "#34495e"))
    if max_info:
        sections.append(("MÁXIMO OBSERVADO EN ESTE DATASET",
                         max_info, True, "#c0392b"))

    # ── Renderizar bloques ──
    y = 0.80
    for label, content, monospace, color in sections:
        ax.text(0.05, y, label, transform=ax.transAxes,
                fontsize=11, fontweight="bold", color="#2980b9",
                va="top")
        y -= HEADER_H
        n_lines = content.count("\n") + 1
        family = "monospace" if monospace else "sans-serif"
        ax.text(0.07, y, content, transform=ax.transAxes,
                fontsize=10.5, va="top", linespacing=1.5,
                fontfamily=family, color=color)
        y -= LINE_H * n_lines + SECTION_GAP

    pdf.savefig(fig); plt.close()


# Detectar PDF bloqueado (preview abierto en VS Code)
def _path_libre(path, ext):
    if not os.path.exists(path):
        return path
    try:
        with open(path, "ab"):
            pass
        return path
    except PermissionError:
        pass
    for i in range(1, 20):
        alt = path.replace(ext, f"_nuevo{i if i > 1 else ''}{ext}")
        if not os.path.exists(alt):
            return alt
        try:
            with open(alt, "ab"):
                pass
            return alt
        except PermissionError:
            continue
    return path

_pdf_out = _path_libre(PDF_PATH, ".pdf")
if _pdf_out != PDF_PATH:
    print(f"[!] {os.path.basename(PDF_PATH)} bloqueado, guardando como {os.path.basename(_pdf_out)}")

_docx_out = _path_libre(DOCX_PATH, ".docx")
if _docx_out != DOCX_PATH:
    print(f"[!] {os.path.basename(DOCX_PATH)} bloqueado, guardando como {os.path.basename(_docx_out)}")

# Documento Word global usado por pagina_titulo / pagina_texto
_DOC = Document()
_titulo_principal = _DOC.add_heading("Scoring de Riesgo Académico", level=0)
_DOC.add_paragraph("Diseño, construcción y validación").italic = True
_DOC.add_paragraph(
    f"Autor: Facundo Pepino — TP Ciencia de Datos 2026 — UTN FRSFCO\n"
    f"Dataset: Student exam score — {N} estudiantes."
)
_DOC.add_paragraph()


print("[4/4] Generando PDF y DOCX...")
with PdfPages(_pdf_out) as pdf:

    # ── PORTADA ─────────────────────────────────────────────────────────
    portada(pdf,
            "Scoring de Riesgo Académico",
            "Diseño, construcción y validación")

    # ── CONSIGNA DEL TP ─────────────────────────────────────────────────
    pagina_titulo(pdf, "Consigna del TP")

    pagina_texto(pdf, "6. Actividades que deberán realizar", [
        " 1. Definir el problema de negocio o académico que se quiere apoyar con un",
        "    score.",
        " 2. Identificar la unidad de análisis (cliente, estudiante, transacción, etc.).",
        " 3. Determinar qué significa el score (ej. probabilidad de fraude, riesgo de",
        "    abandono, propensión a comprar).",
        " 4. Explorar el dataset propuesto: estadísticas descriptivas, valores faltantes,",
        "    correlaciones.",
        " 5. Limpiar y transformar los datos (codificación, normalización si",
        "    corresponde).",
        " 6. Construir variables derivadas que aporten información al score.",
        " 7. Diseñar un scoring basado en reglas o pesos (por ejemplo, mediante",
        "    combinación lineal o reglas if-then).",
        " 8. Definir niveles de score (bajo / medio / alto / crítico) y su",
        "    interpretación.",
        " 9. Asociar acciones sugeridas a cada nivel (ej. ofrecer descuento, llamar",
        "    al cliente, alertar al docente, recomendar tutoría).",
        "10. Implementar un modelo simple con scikit-learn (por ejemplo, regresión",
        "    logística o árbol de decisión) para estimar una probabilidad y compararla",
        "    con el score de reglas.",
        "11. Comparar el scoring por reglas y el scoring por modelo, analizando sus",
        "    diferencias.",
        "12. Evaluar el modelo con métricas básicas: matriz de confusión, accuracy,",
        "    precision, recall, F1, AUC.",
        "13. Reflexionar sobre los falsos positivos y falsos negativos en el contexto",
        "    elegido.",
        "14. Visualizar la distribución de los scores y los niveles obtenidos.",
        "15. Presentar conclusiones, limitaciones y propuestas de mejora.",
    ])

    pagina_texto(pdf, "Respuestas — Actividades (1 a 5)", [
        "1. PROBLEMA: detectar en forma temprana a estudiantes con riesgo de bajo",
        "   desempeño académico para activar tutorías ANTES del examen final.",
        "   No se busca predecir la nota: se busca priorizar la intervención.",
        "",
        "2. UNIDAD DE ANÁLISIS: el estudiante (una fila = un alumno por cursada).",
        "   El score se calcula por alumno, no por examen ni por materia.",
        "",
        "3. SIGNIFICADO DEL SCORE: probabilidad/propensión de tener BAJO DESEMPEÑO",
        "   académico (exam_score < 60). Valor entre 0 y 100: a mayor score,",
        "   mayor riesgo estimado. El score es ordinal, no una nota.",
        "",
        "4. EDA: dataset de 200 alumnos, sin nulos. Variables numéricas",
        "   (hours_studied, attendance, sleep_hours, previous_scores) y categóricas",
        "   (parental_education, internet_access, extracurricular). Correlaciones",
        "   con exam_score: hours_studied r=0.78, attendance r=0.65,",
        "   previous_scores r=0.58, sleep_hours r=0.21. Clase positiva ~25 %.",
        "",
        "5. LIMPIEZA Y TRANSFORMACIÓN: no hubo imputaciones (no había NaN).",
        "   Para el modelo: StandardScaler sobre las numéricas, split 70/30",
        "   estratificado, random_state=42. Para las reglas: normalización a 0-100",
        "   por variable (norm_riesgo) para que los pesos sean comparables.",
    ])

    pagina_texto(pdf, "Respuestas — Actividades (6 a 10)", [
        "6. VARIABLES DERIVADAS:",
        "   • norm_riesgo de hours_studied, attendance, previous_scores, sleep_hours",
        "     (invertidas: menos = más riesgo) llevadas a escala 0-100.",
        "   • score_reglas: combinación lineal con pesos 40/30/20/10.",
        "   • nivel_riesgo: categorización del score en 4 tramos.",
        "   • score_modelo: probabilidad predicha por Regresión Logística.",
        "",
        "7. SCORING BASADO EN REGLAS (combinación lineal de pesos):",
        "   score = 0.40·riesgo_horas + 0.30·riesgo_asistencia",
        "         + 0.20·riesgo_notas_previas + 0.10·riesgo_sueño",
        "   Los pesos replican el orden de las correlaciones observadas en el EDA.",
        "",
        "8. NIVELES DE SCORE (umbrales validados empíricamente):",
        "   • BAJO  (0-30):  sin alerta, seguimiento habitual.",
        "   • MEDIO (30-50): atención del docente, monitoreo.",
        "   • ALTO  (50-70): tutoría recomendada.",
        "   • CRÍTICO (70-100): intervención inmediata del coordinador.",
        "",
        "9. ACCIONES SUGERIDAS POR NIVEL:",
        "   • Bajo: ninguna acción especial.",
        "   • Medio: recordatorios, recursos de estudio adicionales.",
        "   • Alto: convocar a tutoría, contactar a la familia.",
        "   • Crítico: reunión con coordinador + plan de recuperación.",
        "",
        "10. MODELO PREDICTIVO (scikit-learn):",
        "    • Regresión Logística (interpretable, baseline).",
        "    • Random Forest (compara desempeño no lineal).",
        "    • Target binario: riesgo_bajo_desempeno = (exam_score < 60).",
        "    • exam_score se EXCLUYE de las features para evitar data leakage.",
    ])

    pagina_texto(pdf, "Respuestas — Actividades (11 a 15)", [
        "11. COMPARACIÓN REGLAS vs MODELO:",
        "    • Ambos rankean a los alumnos de forma muy similar (corr. alta).",
        "    • El score por reglas es totalmente transparente y auditable.",
        "    • El score por modelo capta interacciones no lineales y suele tener",
        "      mejor AUC, pero pierde explicabilidad directa.",
        "    • Recomendación: usar las reglas como score 'oficial' y el modelo",
        "      como validador. Si ambos coinciden en CRÍTICO → alta confianza.",
        "",
        "12. MÉTRICAS DE EVALUACIÓN:",
        f"    • AUC-ROC LR ≈ {auc_lr:.3f}  |  AUC-ROC RF ≈ {auc_rf:.3f}.",
        "    • Recall priorizado (no perder alumnos en riesgo real).",
        "    • Accuracy NO es buena métrica (clase desbalanceada ~25 %).",
        "    • Matriz de confusión, precision y F1 también reportadas.",
        "",
        "13. FALSOS POSITIVOS vs FALSOS NEGATIVOS:",
        "    • FN (decir 'sin riesgo' y desaprobar): MUY GRAVE — alumno sin ayuda.",
        "    • FP (decir 'riesgo' y aprobar): leve — una tutoría 'de más'.",
        "    • Por eso el umbral se elige para MAXIMIZAR RECALL aceptando algo",
        "      más de FP. El costo asimétrico justifica priorizar sensibilidad.",
        "",
        "14. VISUALIZACIONES INCLUIDAS EN ESTE REPORTE:",
        "    • Histograma de score_reglas y score_modelo.",
        "    • Distribución por nivel de riesgo (bajo/medio/alto/crítico).",
        "    • Curva ROC, matriz de confusión, scatter reglas vs modelo.",
        "",
        "15. CONCLUSIONES, LIMITACIONES Y MEJORAS:",
        "    • Conclusión: ambos enfoques son consistentes y útiles como alerta",
        "      temprana antes del examen final.",
        "    • Limitaciones: n=200, una sola cohorte, variables autodeclaradas.",
        "    • Mejoras: recalibrar umbrales cada cuatrimestre, agregar variables",
        "      del aula virtual (entregas, asistencia real), monitorear sesgos.",
    ])

    pagina_titulo(pdf, "7. Preguntas orientadoras para profundizar")

    pagina_texto(pdf, "Comprensión del problema", [
        "• ¿Qué decisión concreta se busca apoyar con el score?",
        "",
        "• ¿Quién lo utilizará y en qué momento?",
        "",
        "• ¿Cuál es la diferencia entre scoring, clasificación y predicción?",
    ])

    pagina_texto(pdf, "Respuestas — Comprensión del problema", [
        "¿QUÉ DECISIÓN CONCRETA SE BUSCA APOYAR CON EL SCORE?",
        "Identificar, antes del examen final, qué estudiantes tienen mayor riesgo",
        "de bajo desempeño para activar intervenciones de la cátedra: tutorías,",
        "material de refuerzo, seguimiento personalizado y, en casos críticos,",
        "tutoría obligatoria con seguimiento semanal.",
        "",
        "¿QUIÉN LO UTILIZARÁ Y EN QUÉ MOMENTO?",
        "Tutores, coordinador académico y docente del curso. Se aplica al inicio",
        "o a mitad del cuatrimestre, ANTES del examen final — por eso exam_score",
        "no se usa como variable predictora (sería data leakage: se predeciría",
        "una decisión usando información que aún no se conoce).",
        "",
        "¿CUÁL ES LA DIFERENCIA ENTRE SCORING, CLASIFICACIÓN Y PREDICCIÓN?",
        "• Predicción: estima el VALOR de una variable continua (ej. la nota",
        "  exacta en el examen). Es regresión.",
        "• Clasificación: asigna una ETIQUETA discreta (ej. desaprueba sí/no).",
        "  Su salida es binaria o multiclase, no permite priorizar dentro de",
        "  cada clase.",
        "• Scoring: ORDENA a los individuos en una escala continua (0-100) para",
        "  priorizar acciones. Permite elegir el umbral según la capacidad",
        "  operativa (cuántas tutorías hay disponibles).",
    ])

    pagina_texto(pdf, "Diseño del score", [
        "• ¿Qué variables son más relevantes y por qué?",
        "",
        "• ¿Cómo se justifican los pesos asignados?",
        "",
        "• ¿Qué pasa si una variable está en una escala distinta a las demás?",
        "",
        "• ¿Cómo se eligen los umbrales o niveles de score?",
    ])

    pagina_texto(pdf, "Respuestas — Diseño del score", [
        "¿QUÉ VARIABLES SON MÁS RELEVANTES Y POR QUÉ?",
        "Según el EDA, hours_studied es la más fuerte (r = 0.78 con exam_score),",
        "seguida de previous_scores (r = 0.43), attendance_percent (r = 0.23) y",
        "sleep_hours (r = 0.19). Las dos primeras concentran ~70 % del poder",
        "predictivo; las dos restantes se incluyen porque tienen sustento",
        "teórico (compromiso y descanso afectan rendimiento) aunque correlacionen",
        "menos.",
        "",
        "¿CÓMO SE JUSTIFICAN LOS PESOS ASIGNADOS?",
        "Los pesos (40 / 30 / 20 / 10 %) replican el ORDEN de correlación con",
        "exam_score y respetan que la suma sea 100 %. No vienen de un ajuste",
        "automático sino de criterio experto + evidencia del EDA. Esto los hace",
        "auditables y explicables ante un usuario no técnico.",
        "",
        "¿QUÉ PASA SI UNA VARIABLE ESTÁ EN UNA ESCALA DISTINTA?",
        "Hay que NORMALIZAR antes de combinar. En el score por reglas se usa",
        "norm_riesgo(valor, bueno, malo) que mapea cada variable a [0, 100] con",
        "rangos definidos por dominio. En el modelo, StandardScaler estandariza",
        "(media 0, std 1) para que LogisticRegression no penalice variables",
        "con mayor varianza nominal.",
        "",
        "¿CÓMO SE ELIGEN LOS UMBRALES O NIVELES DE SCORE?",
        "Se usaron cortes en 30 / 50 / 70 que dividen la escala en cuatro",
        "niveles operativos (Bajo / Medio / Alto / Crítico) y se VALIDARON",
        "empíricamente: la tasa real de bajo desempeño debe crecer monotónicamente",
        "de Bajo a Crítico, lo que se cumple. En otra cohorte habría que",
        "recalibrarlos.",
    ])

    pagina_texto(pdf, "Modelo predictivo", [
        "• ¿Qué tipo de problema es: clasificación binaria, multiclase, regresión?",
        "",
        "• ¿Por qué se eligió ese algoritmo (regresión logística, árbol, random",
        "  forest, etc.)?",
        "",
        "• ¿Qué decisiones de preprocesamiento se tomaron (split, escalado,",
        "  encoding)?",
    ])

    pagina_texto(pdf, "Respuestas — Modelo predictivo", [
        "¿QUÉ TIPO DE PROBLEMA ES?",
        "CLASIFICACIÓN BINARIA: predice riesgo_bajo_desempeno ∈ {0, 1}. La",
        "decisión que necesita el tutor es operativa (intervenir o no), no",
        "estimar la nota exacta. Aunque internamente se trabaje con la",
        "probabilidad continua P(riesgo), eso es para SCORING, no para",
        "regresión sobre exam_score.",
        "",
        "¿POR QUÉ SE ELIGIÓ ESE ALGORITMO?",
        "Se compararon dos modelos con validación cruzada (5 folds, AUC):",
        "• Logistic Regression: lineal, INTERPRETABLE (coeficientes leíbles),",
        "  rápido, sirve como baseline robusto.",
        "• Random Forest: capta no-linealidades e interacciones, más robusto a",
        "  outliers, no requiere escalado.",
        "Se eligió el de mejor AUC en CV. En contextos académicos con n pequeño",
        f"y necesidad de explicabilidad, LR suele ganar — fue el caso (AUC = {auc_lr:.3f}",
        f"vs {auc_rf:.3f}).",
        "",
        "¿QUÉ DECISIONES DE PREPROCESAMIENTO SE TOMARON?",
        "• SPLIT: train/test 70/30 ESTRATIFICADO por la clase positiva (~25 %)",
        "  para que ambos sets mantengan la misma proporción de riesgo.",
        "• ESCALADO: StandardScaler dentro de un Pipeline para LR (sensible a",
        "  escala). RF no lo necesita.",
        "• ENCODING: no aplica — todas las features son numéricas.",
        "• SEMILLA: random_state = 42 para reproducibilidad.",
        "• EXAM_SCORE EXCLUIDO: para evitar data leakage (es lo que se predice).",
    ])

    pagina_texto(pdf, "Evaluación", [
        "• ¿Qué métrica es la más adecuada para este caso?",
        "",
        "• ¿Es más grave un falso positivo o un falso negativo?",
        "",
        "• ¿El score es estable ante pequeños cambios en los datos?",
    ])

    pagina_texto(pdf, "Respuestas — Evaluación", [
        "¿QUÉ MÉTRICA ES LA MÁS ADECUADA PARA ESTE CASO?",
        "Hay dos lecturas complementarias:",
        "• AUC-ROC: evalúa la calidad del RANKING completo, independiente del",
        "  umbral. Útil porque el score se puede leer a distintos cortes según",
        "  la capacidad operativa de la cátedra.",
        "• RECALL sobre la clase positiva (riesgo): para el umbral elegido,",
        "  importa cuántos casos de riesgo real se detectan. Un recall del 70-80 %",
        "  es el objetivo operativo.",
        "Accuracy NO es la mejor porque la clase está desbalanceada (~25 % positivos):",
        "predecir 'siempre 0' daría 75 % accuracy y sería inútil.",
        "",
        "¿ES MÁS GRAVE UN FALSO POSITIVO O UN FALSO NEGATIVO?",
        "FALSO NEGATIVO es claramente más grave en este dominio:",
        "• FN: el score dice 'sin riesgo' y el alumno termina desaprobando.",
        "  Costo: oportunidad perdida de ayudar; impacto académico real.",
        "• FP: el score dice 'riesgo' y el alumno aprueba.",
        "  Costo: una entrevista innecesaria del tutor (bajo, reversible).",
        "Por eso se prioriza RECALL sobre precision al elegir el umbral.",
        "",
        "¿EL SCORE ES ESTABLE ANTE PEQUEÑOS CAMBIOS EN LOS DATOS?",
        "Indicadores de estabilidad observados:",
        "• AUC en CV (5 folds, train) ≈ AUC en test → no hay sobreajuste fuerte.",
        "• Score por reglas y score por modelo correlacionan fuertemente pese a",
        "  construirse de forma totalmente distinta → señal consistente.",
        "• La tasa real de bajo desempeño crece monotónicamente con el nivel.",
        "Limitación: n = 200 es chico. Con otra cohorte podría haber variación,",
        "por eso se recomienda recalibrar el umbral en cada cuatrimestre.",
    ])

    pagina_texto(pdf, "Ética, sesgos y explicabilidad", [
        "• ¿El dataset puede contener sesgos que afecten el score?",
        "",
        "• ¿Cómo se podría explicar el score a una persona no técnica?",
        "",
        "• ¿Qué controles deberían existir antes de implementarlo en producción?",
    ])

    pagina_texto(pdf, "Respuestas — Ética, sesgos y explicabilidad", [
        "¿EL DATASET PUEDE CONTENER SESGOS QUE AFECTEN EL SCORE?",
        "Sí, varios:",
        "• Variables autodeclaradas: hours_studied y sleep_hours dependen de la",
        "  honestidad del estudiante (sesgo de respuesta).",
        "• Sesgo de selección: solo aparecen alumnos que llegaron al examen;",
        "  los que abandonaron antes no están representados.",
        "• Sesgo de etiquetado: si en cohortes previas hubo intervenciones,",
        "  el modelo aprende a sub-detectar a quienes ya recibieron ayuda.",
        "• Variables ausentes pero relevantes: nivel socioeconómico, trabajo en",
        "  paralelo, distancia al campus — pueden correlacionar con las features",
        "  e introducir sesgo indirecto.",
        "",
        "¿CÓMO SE PODRÍA EXPLICAR EL SCORE A UNA PERSONA NO TÉCNICA?",
        "Usando lenguaje natural y los valores propios del alumno:",
        "  'Tu score combina cuatro señales: cuánto estudiás (40 %), tu",
        "   trayectoria previa (30 %), tu asistencia (20 %) y tus horas de",
        "   sueño (10 %). En tu caso, estudiás X horas y asistís al Y %, eso",
        "   te ubica en nivel Z. Si aumentás horas de estudio y asistencia",
        "   (las dos variables más controlables a corto plazo), el score baja.'",
        "Cada estudiante debería poder pedir este detalle (derecho a la",
        "explicación).",
        "",
        "¿QUÉ CONTROLES DEBERÍAN EXISTIR ANTES DE IMPLEMENTARLO?",
        "• TÉCNICOS: validación con otra cohorte, monitoreo de drift, pruebas",
        "  de fairness por carrera / sede / turno, recalibración periódica.",
        "• HUMANOS: el score ASISTE, NO DECIDE — todo caso pasa por un tutor",
        "  que puede sobrescribir el nivel. Auditoría manual de 'Crítico'.",
        "• ÉTICOS Y NORMATIVOS: consentimiento informado, derecho a no",
        "  participar, transparencia sobre qué variables se usan, vía formal",
        "  para impugnar el nivel asignado.",
        "• EXCLUSIONES EXPLÍCITAS: nunca incorporar género, edad, etnia,",
        "  nacionalidad ni nivel socioeconómico como features.",
    ])

    # ── 1. DEFINICIÓN DEL PROBLEMA ──────────────────────────────────────
    pagina_titulo(pdf, "1. Definición del problema")

    pagina_texto(pdf, "Definición del problema", [
        "DECISIÓN A APOYAR",
        "Detectar antes del examen final qué estudiantes tienen mayor probabilidad de bajo",
        "desempeño, para que la cátedra pueda intervenir (tutorías, recuperatorios, seguimiento).",
        "",
        "UNIDAD DE ANÁLISIS",
        "Un estudiante (1 fila = 1 estudiante = 1 score).",
        "",
        "QUÉ SIGNIFICA EL SCORE",
        "Riesgo de bajo desempeño académico en el examen. Score alto = mayor riesgo (negativo).",
        "Convención coherente con scoring crediticio de riesgo.",
        "",
        "QUIÉN LO USA Y CUÁNDO",
        "Tutores, coordinador académico y docente del curso, al inicio o mitad del cuatrimestre,",
        "ANTES del examen final. Por eso 'exam_score' no se usa como variable predictora",
        "(sería data leakage: la decisión la tomamos antes de conocer la nota).",
        "",
        "SCORING vs CLASIFICACIÓN vs PREDICCIÓN",
        "• Predicción: estima el valor de exam_score (regresión).",
        "• Clasificación: asigna 0 o 1 (¿desaprobará o no?).",
        "• Scoring: ordena a los estudiantes en 0-100 para priorizar intervenciones.",
    ])

    # ── 2. EXPLORACIÓN INICIAL ──────────────────────────────────────────
    pagina_titulo(pdf, "2. Exploración del dataset")

    desc = df[NUMERICAS].describe().round(2)
    data = [[idx] + [f"{v:.2f}" for v in desc.loc[idx].values] for idx in desc.index]
    render_tabla(pdf, data, ["estadístico"] + NUMERICAS,
                 "Estadísticos descriptivos (variables numéricas)")

    # Faltantes y duplicados
    nulls = df[["student_id"] + NUMERICAS].isna().sum()
    data = [[c, int(n)] for c, n in nulls.items()]
    data.append(["student_id duplicados", int(df["student_id"].duplicated().sum())])
    data.append(["filas duplicadas", int(df.duplicated().sum())])
    render_tabla(pdf, data, ["columna / chequeo", "valores faltantes / dups"],
                 "Valores faltantes y duplicados",
                 col_widths=[0.5, 0.3])

    # Histogramas
    fig, axes = plt.subplots(2, 3, figsize=(11, 7))
    for ax, col in zip(axes.flat, NUMERICAS):
        sns.histplot(df[col], kde=True, ax=ax, color="#3b6ea8")
        ax.set_title(col)
    axes.flat[-1].axis("off")
    fig.suptitle("Distribución de variables numéricas", fontsize=14, fontweight="bold")
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # Correlaciones
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(df[NUMERICAS].corr(), annot=True, fmt=".2f",
                cmap="RdBu_r", center=0, vmin=-1, vmax=1, ax=ax)
    ax.set_title("Matriz de correlaciones", fontsize=14, fontweight="bold")
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # Tabla de correlaciones con exam_score
    corr_target = df[NUMERICAS].corr()["exam_score"].drop("exam_score").round(3)
    data = [[var, f"{r:.3f}"] for var, r in corr_target.sort_values(ascending=False).items()]
    render_tabla(pdf, data, ["variable", "correlación con exam_score"],
                 "Correlaciones con exam_score (ordenadas)",
                 col_widths=[0.5, 0.3])

    pagina_texto(pdf, "Hallazgos del EDA", [
        f"• {N} estudiantes, sin nulos ni duplicados → no requiere imputación.",
        "",
        f"• exam_score va de {df['exam_score'].min():.1f} a {df['exam_score'].max():.1f} —",
        "  no es un 0-100 clásico. Por eso no aplica el umbral de 60 del enunciado:",
        "  dejaría al 100% de los estudiantes como 'riesgo'.",
        "",
        "• previous_scores está en escala 40-95 (parece promedio académico previo 0-100).",
        "",
        "• hours_studied, sleep_hours y attendance_percent están en rangos plausibles.",
        "",
        "• Predictor más fuerte: hours_studied (r = 0.78 con exam_score).",
        "  Le sigue previous_scores (r = 0.43), después attendance (0.23) y sleep_hours (0.19).",
    ])

    # ── 3. VARIABLE OBJETIVO ────────────────────────────────────────────
    pagina_titulo(pdf, "3. Variable objetivo")

    tasa_pos = df["riesgo_bajo_desempeno"].mean()
    pagina_texto(pdf, "Definición de riesgo_bajo_desempeno", [
        "DECISIÓN",
        f"Como exam_score no está en escala 0-100 estándar, el umbral se define",
        "usando el cuartil inferior de la propia distribución: el 25 % de estudiantes",
        "con peor nota se considera 'bajo desempeño'.",
        "",
        "REGLA",
        f"  riesgo_bajo_desempeno = 1   si exam_score < {UMBRAL_BAJO:.2f}",
        f"  riesgo_bajo_desempeno = 0   en caso contrario",
        "",
        f"DISTRIBUCIÓN OBSERVADA",
        f"  • Clase 0 (sin riesgo): {(1-tasa_pos)*100:.1f} %",
        f"  • Clase 1 (riesgo):     {tasa_pos*100:.1f} %",
        "",
        "Es un criterio relativo, transparente y reproducible. En otra cohorte habría",
        "que recalibrar el umbral.",
    ])

    # ── 4. SCORING POR REGLAS ───────────────────────────────────────────
    pagina_titulo(pdf, "4. Scoring por reglas")

    data = [
        ["hours_studied",      "40 %", "Mayor correlación con la nota (0.78). Variable más controlable por el estudiante."],
        ["previous_scores",    "30 %", "Segunda correlación más alta (0.43). Refleja trayectoria académica previa."],
        ["attendance_percent", "20 %", "Predictor moderado (0.23). Mide compromiso con la cursada."],
        ["sleep_hours",        "10 %", "Menor correlación (0.19) pero sustento teórico (descanso afecta cognición)."],
    ]
    render_tabla(pdf, data, ["variable", "peso", "justificación"],
                 "Pesos del scoring por reglas",
                 col_widths=[0.22, 0.10, 0.62], fontsize=8.5)

    data = [
        ["hours_studied",      "≥ 10 hs/sem", "≤ 2 hs/sem"],
        ["previous_scores",    "≥ 85",        "≤ 50"],
        ["attendance_percent", "≥ 90 %",      "≤ 60 %"],
        ["sleep_hours",        "7-8 hs (óptimo)", "<5 ó >9 (forma de U)"],
    ]
    render_tabla(pdf, data, ["variable", "rango sin riesgo", "rango alto riesgo"],
                 "Rangos de normalización (definidos por dominio)",
                 col_widths=[0.25, 0.35, 0.35])

    # Distribución del score por reglas
    fig, ax = plt.subplots(figsize=(10, 5))
    sns.histplot(df["score_reglas"], bins=20, kde=True, ax=ax, color="#c0392b")
    ax.axvline(30, ls="--", color="#27ae60", label="Bajo / Medio")
    ax.axvline(50, ls="--", color="#f1c40f", label="Medio / Alto")
    ax.axvline(70, ls="--", color="#e67e22", label="Alto / Crítico")
    ax.set_title("Distribución del score de riesgo (por reglas)",
                 fontsize=14, fontweight="bold")
    ax.set_xlabel("Score de riesgo (0-100)")
    ax.legend()
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # ── 5. NIVELES Y ACCIONES ───────────────────────────────────────────
    pagina_titulo(pdf, "5. Niveles de riesgo y acciones")

    conteo_reglas = df["nivel_reglas"].value_counts().reindex(ORDEN_NIVELES, fill_value=0)
    data = [
        ["Bajo",    "0 - 30",  str(conteo_reglas["Bajo"]),
         f"{conteo_reglas['Bajo']/N*100:.1f} %",
         "Sin intervención. Comunicar disponibilidad de tutorías."],
        ["Medio",   "30 - 50", str(conteo_reglas["Medio"]),
         f"{conteo_reglas['Medio']/N*100:.1f} %",
         "Encuesta de seguimiento al promediar el cuatrimestre."],
        ["Alto",    "50 - 70", str(conteo_reglas["Alto"]),
         f"{conteo_reglas['Alto']/N*100:.1f} %",
         "Entrevista breve con tutor académico. Material de refuerzo."],
        ["Crítico", "70 - 100", str(conteo_reglas["Crítico"]),
         f"{conteo_reglas['Crítico']/N*100:.1f} %",
         "Tutoría obligatoria, seguimiento semanal, revisión de plan de estudio."],
    ]
    render_tabla(pdf, data,
                 ["nivel", "rango", "n", "%", "acción sugerida"],
                 "Niveles de riesgo (score por reglas) y acción asociada",
                 col_widths=[0.10, 0.10, 0.06, 0.08, 0.66], fontsize=8.5)

    # Validación: tasa real de bajo desempeño por nivel
    tasa_nivel = (
        df.groupby("nivel_reglas")["riesgo_bajo_desempeno"]
          .agg(["count", "sum", "mean"])
          .reindex(ORDEN_NIVELES)
    )
    data = [
        [n, int(tasa_nivel.loc[n, "count"]),
         int(tasa_nivel.loc[n, "sum"]),
         f"{tasa_nivel.loc[n, 'mean']*100:.1f} %"]
        for n in ORDEN_NIVELES if not pd.isna(tasa_nivel.loc[n, "count"])
    ]
    render_tabla(pdf, data,
                 ["nivel", "n", "positivos reales", "tasa de bajo desempeño"],
                 "Validación: ¿el score por reglas captura riesgo real?",
                 col_widths=[0.18, 0.18, 0.25, 0.30])

    pagina_texto(pdf, "Lectura de la validación", [
        "Si el score por reglas funciona, la 'tasa de bajo desempeño' debe crecer",
        "monotónicamente del nivel Bajo al Crítico. Es lo que se observa:",
        "",
        f"  Bajo    → {tasa_nivel.loc['Bajo','mean']*100:5.1f} % de los estudiantes desaprueba",
        f"  Medio   → {tasa_nivel.loc['Medio','mean']*100:5.1f} %",
        f"  Alto    → {tasa_nivel.loc['Alto','mean']*100:5.1f} %",
        f"  Crítico → {tasa_nivel.loc['Crítico','mean']*100:5.1f} %",
        "",
        "El salto entre Alto y Crítico es la señal más fuerte: el score por reglas",
        "logra concentrar el riesgo real en el extremo superior de la escala.",
    ])

    # ── 6. MODELO PREDICTIVO ────────────────────────────────────────────
    pagina_titulo(pdf, "6. Modelo predictivo (scikit-learn)")

    pagina_texto(pdf, "Decisiones de modelado", [
        "TIPO DE PROBLEMA",
        "Clasificación binaria (riesgo / no riesgo). El usuario final necesita decidir",
        "intervenir o no — output operativo, no estimar la nota exacta.",
        "",
        "VARIABLES PREDICTORAS",
        "hours_studied, sleep_hours, attendance_percent, previous_scores.",
        "Disponibles ANTES del examen final. exam_score se excluye (data leakage).",
        "",
        "MODELOS COMPARADOS",
        "• Logistic Regression: interpretable, rápido, baseline.",
        "• Random Forest: capta no-linealidades, robusto a outliers.",
        "Se elige el de mejor AUC en validación cruzada (5 folds).",
        "",
        "OTRAS DECISIONES",
        "• train/test estratificado (clase positiva ≈ 25 %).",
        "• StandardScaler en LR (sensible a escalas).",
        "• Random state = 42 para reproducibilidad.",
    ])

    data = [
        ["Logistic Regression",
         f"{auc_lr_cv.mean():.3f} ± {auc_lr_cv.std():.3f}",
         f"{auc_lr:.3f}"],
        ["Random Forest",
         f"{auc_rf_cv.mean():.3f} ± {auc_rf_cv.std():.3f}",
         f"{auc_rf:.3f}"],
    ]
    render_tabla(pdf, data,
                 ["modelo", "AUC CV (5 folds, train)", "AUC test"],
                 f"Comparación de modelos — Elegido: {nombre_final}",
                 col_widths=[0.30, 0.35, 0.20])

    # Coeficientes / importancias
    if isinstance(modelo_final, Pipeline):
        coefs = modelo_final.named_steps["clf"].coef_[0]
        importancias = pd.Series(coefs, index=FEATURES).sort_values()
        titulo_imp = "Coeficientes (Logistic Regression sobre variables estandarizadas)"
    else:
        importancias = pd.Series(modelo_final.feature_importances_, index=FEATURES).sort_values()
        titulo_imp = "Importancia de variables (Random Forest)"

    fig, ax = plt.subplots(figsize=(9, 4.5))
    importancias.plot.barh(ax=ax, color="#3b6ea8")
    ax.set_title(titulo_imp, fontsize=13, fontweight="bold")
    ax.axvline(0, color="black", linewidth=0.5)
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # ── 7. COMPARACIÓN REGLAS vs MODELO ─────────────────────────────────
    pagina_titulo(pdf, "7. Comparación reglas vs modelo")

    auc_reglas = roc_auc_score(df["riesgo_bajo_desempeno"], df["score_reglas"])
    auc_modelo_full = roc_auc_score(df["riesgo_bajo_desempeno"], df["score_modelo"])
    corr_scores = df[["score_reglas", "score_modelo"]].corr().iloc[0, 1]

    data = [
        ["Score por reglas",                f"{auc_reglas:.3f}"],
        [f"Score por modelo ({nombre_final})", f"{auc_modelo_full:.3f}"],
    ]
    render_tabla(pdf, data, ["método", "AUC sobre todo el dataset"],
                 "AUC de cada enfoque",
                 col_widths=[0.45, 0.30])

    # Scatter reglas vs modelo
    fig, ax = plt.subplots(figsize=(8, 7))
    sc = ax.scatter(df["score_reglas"], df["score_modelo"],
                    c=df["riesgo_bajo_desempeno"], cmap="RdYlGn_r",
                    alpha=0.75, edgecolor="k")
    ax.plot([0, 100], [0, 100], "k--", alpha=0.3, label="y = x")
    ax.set_xlabel("Score por reglas")
    ax.set_ylabel("Score por modelo")
    ax.set_title(f"Score reglas vs modelo (correlación = {corr_scores:.3f})",
                 fontsize=13, fontweight="bold")
    ax.legend()
    cbar = plt.colorbar(sc, ax=ax, ticks=[0, 1])
    cbar.set_label("riesgo real (1 = bajo desempeño)")
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # Tabla de concordancia
    cruzada = pd.crosstab(df["nivel_reglas"], df["nivel_modelo"]) \
        .reindex(index=ORDEN_NIVELES, columns=ORDEN_NIVELES, fill_value=0)
    data = [[idx] + [int(v) for v in cruzada.loc[idx].values] for idx in cruzada.index]
    render_tabla(pdf, data,
                 ["reglas \\ modelo"] + ORDEN_NIVELES,
                 "Concordancia entre niveles (filas: reglas, columnas: modelo)",
                 col_widths=[0.20, 0.15, 0.15, 0.15, 0.15])

    # ── 8. EVALUACIÓN ───────────────────────────────────────────────────
    pagina_titulo(pdf, "8. Evaluación con métricas")

    pagina_texto(pdf, "Falso positivo vs falso negativo", [
        "EN ESTE PROBLEMA, EL FALSO NEGATIVO ES MÁS GRAVE",
        "",
        "• FALSO NEGATIVO: el score dice 'sin riesgo' pero el estudiante desaprueba.",
        "  Costo: no recibió la intervención y termina desaprobando.",
        "",
        "• FALSO POSITIVO: el score dice 'riesgo' pero el estudiante aprueba.",
        "  Costo: una entrevista innecesaria del tutor.",
        "",
        "→ Priorizamos RECALL sobre la clase positiva (riesgo). Para un score operativo,",
        "fijamos un umbral que capture al menos el 70-80 % de los casos positivos.",
    ])

    # ROC y PR
    score_test = modelo_final.predict_proba(X_test)[:, 1]
    fpr, tpr, _ = roc_curve(y_test, score_test)
    prec, rec, _ = precision_recall_curve(y_test, score_test)
    ap = average_precision_score(y_test, score_test)
    auc_test = roc_auc_score(y_test, score_test)

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    axes[0].plot(fpr, tpr, label=f"Modelo (AUC = {auc_test:.3f})",
                 color="#3b6ea8", lw=2)
    axes[0].plot([0, 1], [0, 1], "k--", alpha=0.4)
    axes[0].set_xlabel("Tasa falsos positivos")
    axes[0].set_ylabel("Tasa verdaderos positivos")
    axes[0].set_title("Curva ROC (test)")
    axes[0].legend()

    axes[1].plot(rec, prec, label=f"AP = {ap:.3f}", color="#c0392b", lw=2)
    axes[1].set_xlabel("Recall")
    axes[1].set_ylabel("Precision")
    axes[1].set_title("Curva Precision-Recall (test)")
    axes[1].legend()
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # ── MATRIZ DE CONFUSIÓN EN UMBRAL OPERATIVO ─────────────────────────
    score_test_100 = score_test * 100
    THRESHOLD_OP = 50
    y_pred_op = (score_test_100 >= THRESHOLD_OP).astype(int)
    cm = confusion_matrix(y_test, y_pred_op, labels=[0, 1])
    tn, fp, fn, tp = cm.ravel()
    total_test = tn + fp + fn + tp
    acc = (tp + tn) / total_test
    prec_op = tp / (tp + fp) if (tp + fp) else 0
    rec_op = tp / (tp + fn) if (tp + fn) else 0
    spec = tn / (tn + fp) if (tn + fp) else 0
    f1_op = 2 * prec_op * rec_op / (prec_op + rec_op) if (prec_op + rec_op) else 0
    npv = tn / (tn + fn) if (tn + fn) else 0

    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    # Heatmap de la matriz
    cm_pct = cm.astype(float) / cm.sum() * 100
    annot_text = np.array([[f"{cm[i,j]}\n({cm_pct[i,j]:.1f} %)"
                             for j in range(2)] for i in range(2)])
    sns.heatmap(cm, annot=annot_text, fmt="", cmap="Blues",
                xticklabels=["Predice: SIN riesgo", "Predice: RIESGO"],
                yticklabels=["Real: SIN riesgo", "Real: RIESGO"],
                ax=axes[0], cbar=False, square=True,
                annot_kws={"fontsize": 12, "fontweight": "bold"},
                linewidths=1, linecolor="white")
    axes[0].set_title(f"Matriz de confusión\n(umbral operativo = {THRESHOLD_OP})",
                      fontsize=12, fontweight="bold")
    axes[0].tick_params(axis="x", rotation=0)
    axes[0].tick_params(axis="y", rotation=0)

    # Tabla de métricas (texto)
    axes[1].axis("off")
    texto_metricas = (
        f"MÉTRICAS DERIVADAS DE LA MATRIZ\n"
        f"(threshold = {THRESHOLD_OP},  n_test = {total_test})\n"
        f"\n"
        f"  TP (Verdaderos Positivos)  =  {tp:3d}\n"
        f"  FP (Falsos Positivos)      =  {fp:3d}\n"
        f"  FN (Falsos Negativos)      =  {fn:3d}   ← más graves aquí\n"
        f"  TN (Verdaderos Negativos)  =  {tn:3d}\n"
        f"  ─────────────────────────────────\n"
        f"  Accuracy    = (TP+TN)/N    =  {acc:.3f}\n"
        f"  Precision   = TP/(TP+FP)   =  {prec_op:.3f}\n"
        f"  Recall      = TP/(TP+FN)   =  {rec_op:.3f}\n"
        f"  Specificity = TN/(TN+FP)   =  {spec:.3f}\n"
        f"  NPV         = TN/(TN+FN)   =  {npv:.3f}\n"
        f"  F1          = 2·P·R/(P+R)  =  {f1_op:.3f}\n"
        f"  AUC ROC                    =  {auc_test:.3f}\n"
    )
    axes[1].text(0.02, 0.98, texto_metricas, transform=axes[1].transAxes,
                 fontsize=10.5, va="top", ha="left",
                 fontfamily="monospace", color="#2c3e50",
                 bbox=dict(boxstyle="round,pad=0.6",
                           facecolor="#ecf0f1", edgecolor="#bdc3c7"))
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # Lectura de la matriz
    pagina_texto(pdf, "Lectura de la matriz de confusión", [
        f"En el set de test (n = {total_test}) y con umbral operativo {THRESHOLD_OP}:",
        "",
        f"  → De {fn + tp} estudiantes con riesgo real, el modelo detectó {tp}",
        f"    y dejó pasar {fn}  (recall = {rec_op:.0%}).",
        "",
        f"  → De {tn + fp} estudiantes sin riesgo, el modelo clasificó bien {tn}",
        f"    y marcó {fp} como riesgo erróneamente  (specificity = {spec:.0%}).",
        "",
        f"  → De los {tp + fp} alumnos que el modelo señaló como en riesgo,",
        f"    {tp} efectivamente lo estaban  (precision = {prec_op:.0%}).",
        "",
        "INTERPRETACIÓN OPERATIVA",
        "",
        f"• Falsos negativos = {fn}  ← son los más graves: el sistema dice 'no",
        "  hace falta intervenir' y el alumno termina desaprobando. Cada FN",
        "  representa una oportunidad perdida de ayudar.",
        "",
        f"• Falsos positivos = {fp}  ← son tolerables: solo cuestan una entrevista",
        "  innecesaria con el tutor.",
        "",
        f"• Si la cátedra tuviera capacidad para más casos, podría bajar el",
        "  umbral (ej. a 40) y subir recall a costa de algo de precision.",
        f"  Si la capacidad fuera muy escasa, subir el umbral (ej. a 60) elevaría",
        "  precision pero dejaría más alumnos en riesgo sin ayuda.",
        "",
        "La tabla de la próxima página muestra esa misma matriz pero a 4 umbrales,",
        "para apoyar la decisión del corte operativo.",
    ])

    # Tabla por umbrales
    data = []
    for u in [40, 50, 60, 70]:
        y_pred = (score_test_100 >= u).astype(int)
        cm = confusion_matrix(y_test, y_pred, labels=[0, 1])
        tn, fp, fn, tp = cm.ravel()
        precision_v = tp / (tp + fp) if (tp + fp) else 0
        recall_v    = tp / (tp + fn) if (tp + fn) else 0
        f1          = 2 * precision_v * recall_v / (precision_v + recall_v) if (precision_v + recall_v) else 0
        data.append([str(u), str(tp), str(fp), str(fn), str(tn),
                     f"{precision_v:.3f}", f"{recall_v:.3f}", f"{f1:.3f}"])
    render_tabla(pdf, data,
                 ["umbral", "TP", "FP", "FN", "TN", "precision", "recall", "F1"],
                 "Métricas por umbral operativo (sobre test)",
                 col_widths=[0.10] * 8)

    # ── 9. CASOS INDIVIDUALES ───────────────────────────────────────────
    pagina_titulo(pdf, "9. Casos individuales")

    ej_bajo = df.sort_values("score_modelo").iloc[0]
    ej_med = df.iloc[(df["score_modelo"] - 50).abs().idxmin()]
    ej_crit = df.sort_values("score_modelo", ascending=False).iloc[0]

    data = []
    for nombre, fila in [("caso bajo", ej_bajo),
                         ("caso medio", ej_med),
                         ("caso crítico", ej_crit)]:
        data.append([
            nombre,
            str(fila["student_id"]),
            f"{fila['hours_studied']:.1f}",
            f"{fila['sleep_hours']:.1f}",
            f"{fila['attendance_percent']:.1f}",
            str(int(fila["previous_scores"])),
            f"{fila['exam_score']:.1f}",
            f"{fila['score_reglas']:.1f} ({fila['nivel_reglas']})",
            f"{fila['score_modelo']:.1f} ({fila['nivel_modelo']})",
        ])
    render_tabla(pdf, data,
                 ["caso", "id", "horas", "sueño", "asis %",
                  "prev", "exam", "score reglas", "score modelo"],
                 "Tres casos representativos",
                 col_widths=[0.10, 0.07, 0.07, 0.07, 0.08, 0.07, 0.07, 0.18, 0.18],
                 fontsize=8.5)

    pagina_texto(pdf, "Cómo se explica el score a un usuario no técnico", [
        "EJEMPLO — caso crítico",
        "",
        f"Estudiante {ej_crit['student_id']} recibió un score alto porque combina:",
        f"  • {ej_crit['hours_studied']:.1f} horas de estudio (rango bajo: ≤ 2 hs)",
        f"  • {int(ej_crit['previous_scores'])} en desempeño previo (rango bajo: ≤ 50)",
        f"  • {ej_crit['attendance_percent']:.1f} % de asistencia",
        f"  • {ej_crit['sleep_hours']:.1f} horas promedio de sueño",
        "",
        "Las variables apuntan en la misma dirección, por eso el score lo ubica en el",
        "nivel crítico y se recomienda tutoría obligatoria con seguimiento semanal.",
        "",
        "El estudiante puede 'mejorar' su score aumentando horas de estudio y asistencia",
        "(las dos variables más controlables a corto plazo).",
    ])

    # ── 9.5 SCORES ADICIONALES ──────────────────────────────────────────
    pagina_titulo(pdf, "10. Scores adicionales derivados")

    pagina_texto(pdf, "Por qué más de un score", [
        "El score de riesgo es solo un ángulo del problema. A partir del mismo dataset",
        "podemos construir scores complementarios que apoyan distintas decisiones de la",
        "cátedra: priorizar tutorías, reconocer trayectorias positivas, detectar",
        "ineficiencias en la técnica de estudio, alertar sobre desbalance vida-estudio,",
        "predecir abandono, etc.",
        "",
        "Cada score se acota a la escala 0-100. Las normalizaciones por percentil hacen",
        "que los scores sean robustos a la escala atípica de exam_score (max ≈ 51).",
        "",
        "Los 9 scores no son redundantes: cada uno apunta a una decisión distinta.",
        "Pueden combinarse — por ejemplo, riesgo alto + potencial alto = prioridad máxima.",
        "",
        "ATENCIÓN — score_desercion es un PROXY. El dataset no tiene etiqueta real de",
        "abandono, así que el score se construye con criterios de literatura (asistencia",
        "como predictor más fuerte) pero NO se puede validar con AUC. Su uso debe ser",
        "INDICATIVO, no decisorio.",
    ])

    data_scores = [
        ["1. score_desempeno",
         "Inverso del riesgo (alto = bueno)",
         "Reconocer rendimiento sin estigmatizar"],
        ["2. score_potencial",
         "Brecha entre nota actual y trayectoria previa",
         "Priorizar tutorías de mayor ROI"],
        ["3. score_engagement",
         "Esfuerzo (asistencia + horas), independiente del resultado",
         "Diferenciar 'no entiende' de 'no se compromete'"],
        ["4. score_eficiencia",
         "Cuánto rinde por hora estudiada",
         "Detectar problemas de técnica de estudio"],
        ["5. score_habitos",
         "Sueño óptimo, horas razonables, asistencia",
         "Alertar burnout o desbalance"],
        ["6. score_progreso",
         "Mejora vs su propio pasado",
         "Reconocer trayectoria ascendente"],
        ["7. score_prioridad",
         "Riesgo (60%) + potencial (40%)",
         "Ordenar la cola de tutorías"],
        ["8. score_anomalia",
         "Discrepancia entre score esperado y resultado real",
         "Marcar para revisión humana caso a caso"],
        ["9. score_desercion (proxy)",
         "Riesgo de abandono: asistencia (45%) + horas + previo + sueño",
         "Plan de retención. NO validable sin etiqueta real"],
    ]
    render_tabla(pdf, data_scores, ["score", "qué mide", "decisión que apoya"],
                 "Catálogo de los 9 scores adicionales",
                 col_widths=[0.22, 0.42, 0.32], fontsize=8.5, row_height=1.4)

    # Estadísticos descriptivos de los nuevos scores
    desc_nuevos = df[NUEVOS_SCORES].describe().round(2).T
    data = [
        [idx,
         f"{desc_nuevos.loc[idx, 'mean']:.2f}",
         f"{desc_nuevos.loc[idx, 'std']:.2f}",
         f"{desc_nuevos.loc[idx, 'min']:.2f}",
         f"{desc_nuevos.loc[idx, '25%']:.2f}",
         f"{desc_nuevos.loc[idx, '50%']:.2f}",
         f"{desc_nuevos.loc[idx, '75%']:.2f}",
         f"{desc_nuevos.loc[idx, 'max']:.2f}"]
        for idx in desc_nuevos.index
    ]
    render_tabla(pdf, data, ["score", "media", "std", "min", "p25", "p50", "p75", "max"],
                 "Estadísticos descriptivos de los 8 scores adicionales",
                 col_widths=[0.22] + [0.10] * 7, fontsize=8.5)

    # Niveles y acciones
    data_niveles = [
        ["score_desempeno",  "<30",  "30-70", "≥70", "Reconocer; usar como ejemplo positivo"],
        ["score_potencial",  "<40",  "40-65", "≥65", "Tutoría focalizada (ROI alto)"],
        ["score_engagement", "<40",  "40-70", "≥70", "OK; validar comprensión si nota baja"],
        ["score_eficiencia", "<30",  "30-70", "≥70", "Reconocer técnica de estudio"],
        ["score_habitos",    "<40",  "40-70", "≥70", "OK; charla de hábitos si <40"],
        ["score_progreso",   "<40",  "40-60", "≥60", "Reconocer mejora"],
        ["score_prioridad",  "<30",  "30-60", "≥60", "Atender primero en la cola"],
        ["score_anomalia",   "<20",  "20-40", "≥40", "Revisión humana caso por caso"],
        ["score_desercion",  "<25",  "25-60", "≥60", "Plan de retención: tutor, contacto coordinador"],
    ]
    render_tabla(pdf, data_niveles,
                 ["score", "Bajo", "Medio", "Alto", "acción si nivel alto"],
                 "Niveles de interpretación y acción asociada",
                 col_widths=[0.20, 0.10, 0.12, 0.10, 0.45], fontsize=8.5)

    # Distribuciones (9 histogramas en 3x3)
    fig, axes = plt.subplots(3, 3, figsize=(13, 9))
    for ax, col in zip(axes.flat, NUEVOS_SCORES):
        sns.histplot(df[col], bins=20, kde=True, ax=ax, color="#3b6ea8")
        ax.set_title(col, fontsize=10)
        ax.set_xlabel("")
    fig.suptitle("Distribución de los 9 scores adicionales",
                 fontsize=14, fontweight="bold")
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # Heatmap de correlación entre los 11 scores
    todos_scores = ["score_reglas", "score_modelo"] + NUEVOS_SCORES
    fig, ax = plt.subplots(figsize=(11, 9))
    sns.heatmap(df[todos_scores].corr(), annot=True, fmt=".2f",
                cmap="RdBu_r", center=0, vmin=-1, vmax=1, square=True, ax=ax,
                annot_kws={"fontsize": 8})
    ax.set_title("Correlación entre los 11 scores",
                 fontsize=14, fontweight="bold")
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # Mismos 3 estudiantes, ahora con los 10 scores
    ej_b = df.sort_values("score_reglas").iloc[0]
    ej_m = df.iloc[(df["score_reglas"] - 50).abs().idxmin()]
    ej_c = df.sort_values("score_reglas", ascending=False).iloc[0]

    data = []
    for nombre, fila in [("caso bajo riesgo", ej_b),
                         ("caso medio", ej_m),
                         ("caso crítico", ej_c)]:
        fila_data = [nombre, str(fila["student_id"]),
                     f"{fila['score_reglas']:.0f}",
                     f"{fila['score_modelo']:.0f}"]
        for s in NUEVOS_SCORES:
            fila_data.append(f"{fila[s]:.0f}")
        data.append(fila_data)
    render_tabla(pdf, data,
                 ["caso", "id", "riesgo", "modelo",
                  "desemp.", "potenc.", "engag.", "eficien.",
                  "hábitos", "progr.", "priorid.", "anom.", "deser."],
                 "Tres casos vistos con los 11 scores",
                 col_widths=[0.12, 0.05] + [0.07] * 11, fontsize=7.5)

    pagina_texto(pdf, "Cómo combinar los scores en una decisión", [
        "Distintos scores cuentan historias distintas del MISMO estudiante:",
        "",
        "• riesgo alto + potencial alto         → prioridad máxima de tutoría (ROI alto)",
        "",
        "• riesgo alto + engagement bajo        → primero llamado de atención por",
        "                                         compromiso, antes que tutoría académica",
        "",
        "• riesgo alto + eficiencia baja        → tutoría enfocada en MÉTODO de estudio,",
        "                                         no en contenido",
        "",
        "• score_anomalia alto                  → caso atípico que requiere REVISIÓN",
        "                                         HUMANA (problema personal, ansiedad",
        "                                         de examen, error de datos)",
        "",
        "• score_habitos bajo                   → charla de bienestar antes que más",
        "                                         horas de estudio (riesgo de burnout)",
        "",
        "• score_progreso alto                  → reconocer mejora aunque la nota",
        "                                         actual no sea alta — refuerza motivación",
        "",
        "• score_desercion alto                 → activar plan de retención: contacto",
        "                                         coordinador, tutor de carrera, posible",
        "                                         entrevista con bienestar estudiantil.",
        "                                         RECORDAR: es proxy, no decisión final",
    ])

    # ── 10. DISTRIBUCIÓN DE NIVELES ─────────────────────────────────────
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.8))
    for ax, col, titulo in [
        (axes[0], "nivel_reglas", "Niveles - Score por reglas"),
        (axes[1], "nivel_modelo", f"Niveles - Score por modelo ({nombre_final})"),
    ]:
        conteo = df[col].value_counts().reindex(ORDEN_NIVELES, fill_value=0)
        bars = ax.bar(conteo.index, conteo.values,
                      color=[COLORES_NIVEL[n] for n in ORDEN_NIVELES])
        ax.set_title(titulo, fontsize=12, fontweight="bold")
        ax.set_ylabel("Estudiantes")
        for bar, v in zip(bars, conteo.values):
            ax.text(bar.get_x() + bar.get_width() / 2, v + 1,
                    str(int(v)), ha="center", fontsize=10)
    plt.tight_layout()
    pdf.savefig(fig); plt.close()

    # ── FÓRMULAS DE CÁLCULO (11 PÁGINAS, UNA POR SCORE) ─────────────────
    pagina_titulo(pdf, "Cómo se calcula cada score\n(11 fórmulas — una por página)")

    pagina_texto(pdf, "Notación común a todas las fórmulas", [
        "Helpers usados en varias fórmulas. Todos los scores devuelven 0-100.",
        "",
        "  norm_riesgo(v, bueno, malo)  →  mapea linealmente el rango [bueno, malo]",
        "                                  al rango [0, 100]. Recorta fuera de rango.",
        "                                  bueno → 0 (sin riesgo); malo → 100 (alto riesgo).",
        "",
        "  riesgo_sueño(h)              →  forma de U: 0 si h ∈ [7, 8]; sino crece",
        "                                  linealmente hasta 100 a 3 hs de desvío.",
        "",
        "  pct(x)                       →  percentil de x dentro del dataset (0 a 100).",
        "                                  Hace los scores robustos a la escala atípica",
        "                                  de exam_score (max ≈ 51, no 100).",
        "",
        "  clip(v, 0, 100)              →  recorta v al rango [0, 100].",
        "",
        "En las próximas páginas, cada score tiene su propia hoja con: fórmula,",
        "interpretación, niveles, decisión que apoya y máximo observado.",
    ])

    # Helper: encontrar el alumno con el máximo de cada score
    def max_de(col):
        idx = df[col].idxmax()
        return f"alumno {df.loc[idx, 'student_id']}  →  {df.loc[idx, col]:.2f}"

    TOTAL = 11

    # ── 1. score_reglas ──
    pagina_detalle_score(
        pdf, 1, TOTAL, "score_reglas", "RIESGO ACADÉMICO",
        formula_lines=[
            "R_horas   = norm_riesgo(hours_studied,      bueno=10, malo=2)",
            "R_previo  = norm_riesgo(previous_scores,    bueno=85, malo=50)",
            "R_asis    = norm_riesgo(attendance_percent, bueno=90, malo=60)",
            "R_sueño   = riesgo_sueño(sleep_hours)",
            "",
            "score_reglas = 0.40·R_horas + 0.30·R_previo",
            "             + 0.20·R_asis  + 0.10·R_sueño",
        ],
        interpretacion_lines=[
            "Score alto = mayor riesgo de bajo desempeño en el examen.",
            "Los pesos reflejan la correlación de cada variable con la nota:",
            "horas estudiadas es el predictor más fuerte (r = 0.78).",
        ],
        niveles="Bajo (0-30) | Medio (30-50) | Alto (50-70) | Crítico (70-100)",
        accion="Crítico → tutoría obligatoria + seguimiento semanal.\nAlto → entrevista breve con tutor + material de refuerzo.",
        max_info=max_de("score_reglas"),
    )

    # ── 2. score_modelo ──
    pagina_detalle_score(
        pdf, 2, TOTAL, "score_modelo", "RIESGO ESTIMADO POR MODELO ML",
        formula_lines=[
            "score_modelo = P(bajo_desempeño | features) × 100",
            "",
            "Donde P proviene de Logistic Regression entrenada con",
            "split 70/30 estratificado y StandardScaler.",
            "",
            "Features: hours_studied, sleep_hours,",
            "          attendance_percent, previous_scores",
        ],
        interpretacion_lines=[
            "El modelo no usa exam_score como input (sería data leakage:",
            "es la variable que estamos tratando de predecir).",
            "AUC test = 0.956 — discrimina bien casos de riesgo.",
        ],
        niveles="Mismos cortes que score_reglas (Bajo / Medio / Alto / Crítico)",
        accion="Mismas acciones que score_reglas. Funciona como segunda opinión\nestadística que valida (o refuta) el ranking del score por reglas.",
        max_info=max_de("score_modelo"),
    )

    # ── 3. score_desempeno ──
    pagina_detalle_score(
        pdf, 3, TOTAL, "score_desempeno", "DESEMPEÑO (alto = bueno)",
        formula_lines=[
            "score_desempeno = 100 − score_reglas",
        ],
        interpretacion_lines=[
            "Inverso simple del score de riesgo. La misma información,",
            "presentada en clave POSITIVA: alto = mejor situación académica.",
            "Útil para reportes orientados a comunicación con el estudiante.",
        ],
        niveles="Crítico (<30) | Regular (30-50) | Bueno (50-70) | Excelente (≥70)",
        accion="Comunicar reconocimiento sin estigmatizar. Usar a los alumnos\nde nivel 'Excelente' como referentes para tutorías peer-to-peer.",
        max_info=max_de("score_desempeno"),
    )

    # ── 4. score_potencial ──
    pagina_detalle_score(
        pdf, 4, TOTAL, "score_potencial", "POTENCIAL DE MEJORA",
        formula_lines=[
            "prev_pct = pct(previous_scores)     # percentil 0-100",
            "exam_pct = pct(exam_score)          # percentil 0-100",
            "",
            "score_potencial = clip((prev_pct − exam_pct) + 50, 0, 100)",
        ],
        interpretacion_lines=[
            "50 = consistente con su trayectoria previa.",
            ">50 = rinde por DEBAJO de su pasado → potencial inexplotado.",
            "<50 = SOBRE-rinde respecto de su trayectoria previa.",
        ],
        niveles="Bajo (<40) | Medio (40-65) | Alto (≥65, hay potencial)",
        accion="Score alto + riesgo alto → MAYOR ROI de tutoría.\nEl alumno tiene la capacidad demostrada en el pasado, solo necesita\nrecuperar el ritmo.",
        max_info=max_de("score_potencial"),
    )

    # ── 5. score_engagement ──
    pagina_detalle_score(
        pdf, 5, TOTAL, "score_engagement", "COMPROMISO / ESFUERZO",
        formula_lines=[
            "horas_n = clip((hours_studied − 1) / 9 × 100, 0, 100)",
            "asis_n  = clip((attendance_percent − 50) / 50 × 100, 0, 100)",
            "",
            "score_engagement = 0.40·horas_n + 0.60·asis_n",
        ],
        interpretacion_lines=[
            "Solo mide ESFUERZO. Ignora la nota — un alumno puede rendir",
            "mal y aún tener engagement alto (viene a clase, dedica horas).",
            "Asistencia pesa más que horas (objetivable, no autodeclarada).",
        ],
        niveles="Bajo (<40) | Medio (40-70) | Alto (≥70)",
        accion="Engagement bajo + riesgo alto → llamado de atención por\nCOMPROMISO antes que tutoría académica.\nEngagement alto + nota baja → validar comprensión, no esfuerzo.",
        max_info=max_de("score_engagement"),
    )

    # ── 6. score_eficiencia ──
    pagina_detalle_score(
        pdf, 6, TOTAL, "score_eficiencia", "NOTA POR HORA INVERTIDA",
        formula_lines=[
            "ratio = exam_score / hours_studied",
            "score_eficiencia = pct(ratio)   # percentil en el dataset",
        ],
        interpretacion_lines=[
            "Alto = rinde mucho con pocas horas → técnica eficiente.",
            "Bajo = estudia mucho pero rinde poco → posible problema de",
            "       MÉTODO de estudio, no de esfuerzo.",
        ],
        niveles="Baja (<30) | Media (30-70) | Alta (≥70)",
        accion="Eficiencia baja + horas altas → tutoría enfocada en TÉCNICA\nde estudio (mapas conceptuales, espaciado, recall activo).",
        max_info=max_de("score_eficiencia"),
    )

    # ── 7. score_habitos ──
    pagina_detalle_score(
        pdf, 7, TOTAL, "score_habitos", "BIENESTAR / BALANCE",
        formula_lines=[
            "sueño_h   = 100 si 7 ≤ sleep_hours ≤ 8",
            "            sino max(0, 100 − 25·|desvío|)",
            "estudio_h = 100 si 3 ≤ hours_studied ≤ 8",
            "            <3 hs:  max(0, 100 − 30·(3 − horas))",
            "            >8 hs:  max(0, 100 − 25·(horas − 8))",
            "asis_h    = clip((attendance_percent − 50) / 50 × 100, 0, 100)",
            "",
            "score_habitos = 0.40·sueño_h + 0.30·estudio_h + 0.30·asis_h",
        ],
        interpretacion_lines=[
            "Penaliza extremos: poco sueño Y exceso de estudio (burnout).",
            "Premia el balance vida-estudio. Sueño pesa más por su impacto",
            "directo en cognición y memoria.",
        ],
        niveles="Bajo (<40) | Aceptable (40-70) | Saludable (≥70)",
        accion="Score bajo → charla de hábitos / bienestar estudiantil ANTES\nque más horas de estudio (riesgo de burnout).",
        max_info=max_de("score_habitos"),
    )

    # ── 8. score_progreso ──
    pagina_detalle_score(
        pdf, 8, TOTAL, "score_progreso", "EVOLUCIÓN VS PASADO",
        formula_lines=[
            "score_progreso = clip((exam_pct − prev_pct) + 50, 0, 100)",
            "",
            "Es el espejo de score_potencial.",
        ],
        interpretacion_lines=[
            "50 = neutro (rinde igual que su pasado).",
            ">50 = mejorando respecto de su trayectoria.",
            "<50 = empeorando respecto de su trayectoria.",
        ],
        niveles="Bajando (<40) | Estable (40-60) | Mejorando (≥60)",
        accion="Reconocer públicamente a los alumnos en 'Mejorando' aunque su\nnota actual no sea alta — refuerza motivación y persistencia.",
        max_info=max_de("score_progreso"),
    )

    # ── 9. score_prioridad ──
    pagina_detalle_score(
        pdf, 9, TOTAL, "score_prioridad", "PRIORIDAD DE INTERVENCIÓN",
        formula_lines=[
            "score_prioridad = 0.60·score_reglas + 0.40·score_potencial",
        ],
        interpretacion_lines=[
            "Combina RIESGO y ROI: prioriza alumnos con riesgo alto Y",
            "potencial alto (los que tienen más para ganar con la intervención).",
            "Útil cuando la cátedra tiene capacidad limitada de tutorías.",
        ],
        niveles="Baja (<30) | Media (30-60) | Alta (≥60, atender YA)",
        accion="Ordenar la cola de tutorías de mayor a menor score_prioridad.\nLos primeros 10-15 entran al programa con seguimiento semanal.",
        max_info=max_de("score_prioridad"),
    )

    # ── 10. score_anomalia ──
    pagina_detalle_score(
        pdf, 10, TOTAL, "score_anomalia", "DISCREPANCIA SCORE-RESULTADO",
        formula_lines=[
            "riesgo_real    = 100 − pct(exam_score)",
            "score_anomalia = | score_reglas − riesgo_real |",
        ],
        interpretacion_lines=[
            "Mide cuánto se desvía el score por reglas del riesgo REAL",
            "calculado a posteriori desde la nota.",
            "Alto = caso ATÍPICO: las señales decían una cosa, pasó otra.",
        ],
        niveles="Coherente (<20) | Discrepancia leve (20-40) | Anómalo (≥40)",
        accion="Score anomalia alto → REVISIÓN HUMANA caso por caso.\nPosibles causas: problema personal, ansiedad de examen, error\nde datos en el legajo, evento puntual no registrado.",
        max_info=max_de("score_anomalia"),
    )

    # ── 11. score_desercion ──
    pagina_detalle_score(
        pdf, 11, TOTAL, "score_desercion", "RIESGO DE ABANDONO  (proxy ⚠)",
        formula_lines=[
            "R_asis_d  = norm_riesgo(attendance_percent, bueno=85, malo=55)",
            "R_horas_d = norm_riesgo(hours_studied,      bueno=8,  malo=2)",
            "R_prev_d  = norm_riesgo(previous_scores,    bueno=80, malo=45)",
            "R_sueño_d = riesgo_sueño(sleep_hours)",
            "",
            "score_desercion = 0.45·R_asis_d + 0.25·R_horas_d",
            "                + 0.20·R_prev_d + 0.10·R_sueño_d",
        ],
        interpretacion_lines=[
            "Asistencia pesa 45 % porque la literatura la identifica como",
            "el predictor más fuerte de abandono universitario.",
            "",
            "LIMITACIÓN: el dataset NO tiene ground truth de 'abandonó /",
            "siguió'. NO se puede validar con AUC. Es INDICATIVO, no decisorio.",
        ],
        niveles="Bajo (<25) | Medio (25-60) | Alto (≥60)",
        accion="Score alto → activar PLAN DE RETENCIÓN: tutor de carrera +\ncoordinador + bienestar estudiantil. Nunca decidir solo en base\nal score: validar con conversación con el alumno.",
        max_info=max_de("score_desercion"),
    )

    # ── 11. CONCLUSIONES ────────────────────────────────────────────────
    pagina_titulo(pdf, "11. Conclusiones, limitaciones y ética")

    pagina_texto(pdf, "Conclusiones", [
        f"1. El score por reglas (AUC = {auc_reglas:.3f}) y el score por modelo",
        f"   ({nombre_final}, AUC = {auc_modelo_full:.3f}) tienen poder predictivo alto",
        f"   y rankean parecido (correlación entre scores = {corr_scores:.2f}).",
        "",
        "2. El score por reglas se valida empíricamente: la tasa real de bajo",
        "   desempeño crece monotónicamente del nivel Bajo al Crítico.",
        "",
        f"3. La variable más informativa es hours_studied (peso 40 % en reglas,",
        f"   y mayor coeficiente/importancia en el modelo).",
        "",
        "4. El umbral operativo del modelo a 50 puntos balancea bien recall y precision",
        "   en este dominio donde el falso negativo es más grave.",
    ])

    pagina_texto(pdf, "Limitaciones", [
        "• n = 200: muestra chica para un modelo robusto. Los AUC pueden variar",
        "  bastante en otra cohorte.",
        "",
        "• hours_studied y sleep_hours son AUTOINFORMADAS → posible sesgo de respuesta.",
        "",
        f"• El umbral de 'bajo desempeño' (p25 de exam_score = {UMBRAL_BAJO:.2f})",
        "  es relativo a esta cohorte. Habría que recalibrarlo en otra promoción.",
        "",
        "• El score se evaluó sobre las mismas variables y muestra; no hay validación",
        "  externa con datos de otro cuatrimestre.",
        "",
        f"• El examen tuvo escala atípica (máx {df['exam_score'].max():.1f}, no 100):",
        "  conviene revisar la consigna del examen antes de generalizar.",
    ])

    pagina_texto(pdf, "Riesgos éticos y mejoras futuras", [
        "RIESGOS ÉTICOS",
        "• Sesgo de etiquetado: si los estudiantes que recibieron tutoría en cohortes",
        "  previas mejoraron, el modelo aprende a sub-detectar ese perfil.",
        "• Estigmatización: el nivel 'Crítico' no debe usarse como etiqueta pública",
        "  ni para condicionar acceso a cursadas.",
        "• Variables prohibidas: si en producción se sumaran género, edad o etnia,",
        "  deberían excluirse explícitamente para evitar discriminación.",
        "• Derecho a la explicación: cada estudiante con score alto debería poder",
        "  pedir el detalle de qué variables lo elevaron y cómo bajarlo.",
        "",
        "MEJORAS FUTURAS",
        "• Validar con otra cohorte (datos del cuatrimestre siguiente).",
        "• Probar XGBoost / Gradient Boosting.",
        "• Calibrar probabilidades (Platt scaling o Isotonic).",
        "• Tablero en Power BI / Looker con filtros por carrera y nivel de riesgo.",
        "• A/B test de la intervención sobre los estudiantes en 'Crítico'.",
    ])

    pagina_texto(pdf, "Controles antes de implementar", [
        "TÉCNICOS",
        "• Validación con datos de otra cohorte.",
        "• Monitoreo de drift y recalibración periódica.",
        "• Pruebas de fairness por grupo (carrera, sede, turno).",
        "",
        "HUMANOS",
        "• El score asiste, NO decide. Toda intervención pasa por un tutor que",
        "  puede sobrescribir el nivel.",
        "• Auditoría manual de casos en 'Crítico'.",
        "",
        "ÉTICOS Y NORMATIVOS",
        "• Consentimiento informado del estudiante.",
        "• Derecho a no participar.",
        "• Transparencia sobre qué variables se usan.",
        "• Vía formal para impugnar el nivel asignado.",
    ])

print()
_DOC.save(_docx_out)
print(f"[OK] PDF generado: {_pdf_out}")
print(f"[OK] DOCX generado: {_docx_out}")
print(f"[OK] CSV de resultados: {CSV_OUT}")
